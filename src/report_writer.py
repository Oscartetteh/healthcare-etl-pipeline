from __future__ import annotations

import os
from pathlib import Path
from typing import Any
from uuid import uuid4

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import PROC_DATA_PATH, RAW_DATA_PATH


def export_report_docx(pipeline: Any, transform_summary: dict, output_path: Path) -> Path:
    """
    Write the ETL execution report to a Microsoft Word document.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    title = document.add_heading("Module 05 - ETL Pipeline Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"Industry: {pipeline.industry}")
    document.add_paragraph(f"Final status: {pipeline._status.upper()}")

    _add_key_value_table(
        document,
        "Extraction",
        [
            ("Source file", RAW_DATA_PATH.name),
            ("Rows loaded", f"{len(pipeline.raw_df):,}"),
            ("Columns loaded", str(pipeline.raw_df.shape[1])),
        ],
    )

    _add_key_value_table(
        document,
        "Validation",
        [
            ("Result", "PASSED" if pipeline.validator._passed else "FAILED"),
            ("Issues found", str(len(pipeline.validator.issues))),
            (
                "Critical",
                str(sum(1 for i in pipeline.validator.issues if i["severity"] == "CRITICAL")),
            ),
            (
                "Warnings",
                str(sum(1 for i in pipeline.validator.issues if i["severity"] == "WARNING")),
            ),
        ],
    )

    _add_issues_table(document, pipeline.validator.issues)

    _add_key_value_table(
        document,
        "Transformation",
        [
            ("Rows before", f"{transform_summary['original_rows']:,}"),
            ("Rows after", f"{transform_summary['final_rows']:,}"),
            ("Rows removed", f"{transform_summary['rows_removed']:,}"),
            ("Columns now", str(transform_summary["final_columns"])),
            ("Changes made", str(transform_summary["changes_count"])),
        ],
    )

    _add_list(document, "Transformation Change Log", transform_summary["change_log"])

    file_size_kb = PROC_DATA_PATH.stat().st_size / 1024
    _add_key_value_table(
        document,
        "Load",
        [
            ("Output file", str(PROC_DATA_PATH)),
            ("File size", f"{file_size_kb:.1f} KB"),
        ],
    )

    _add_list(document, "Pipeline Audit Log", pipeline._run_log)

    document.save(output_path)
    return output_path


def export_report_xlsx(pipeline: Any, transform_summary: dict, output_path: Path) -> Path:
    """
    Write the ETL execution report to a Microsoft Excel workbook.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    file_size_kb = PROC_DATA_PATH.stat().st_size / 1024
    summary_rows = [
        ("Industry", pipeline.industry),
        ("Final status", pipeline._status.upper()),
        ("Source file", RAW_DATA_PATH.name),
        ("Rows loaded", len(pipeline.raw_df)),
        ("Columns loaded", pipeline.raw_df.shape[1]),
        ("Validation result", "PASSED" if pipeline.validator._passed else "FAILED"),
        ("Issues found", len(pipeline.validator.issues)),
        ("Critical issues", sum(1 for i in pipeline.validator.issues if i["severity"] == "CRITICAL")),
        ("Warnings", sum(1 for i in pipeline.validator.issues if i["severity"] == "WARNING")),
        ("Rows before transformation", transform_summary["original_rows"]),
        ("Rows after transformation", transform_summary["final_rows"]),
        ("Rows removed", transform_summary["rows_removed"]),
        ("Columns after transformation", transform_summary["final_columns"]),
        ("Changes made", transform_summary["changes_count"]),
        ("Processed output file", str(PROC_DATA_PATH)),
        ("Processed output size KB", round(file_size_kb, 1)),
    ]

    issues_df = pd.DataFrame(pipeline.validator.issues)
    if issues_df.empty:
        issues_df = pd.DataFrame(columns=["severity", "column", "message"])

    changes_df = pd.DataFrame(
        {"step": range(1, len(transform_summary["change_log"]) + 1), "change": transform_summary["change_log"]}
    )
    audit_df = pd.DataFrame(
        {"step": range(1, len(pipeline._run_log) + 1), "entry": pipeline._run_log}
    )

    temp_path = output_path.with_name(
        f".{output_path.stem}.{uuid4().hex}.tmp{output_path.suffix}"
    )

    try:
        with pd.ExcelWriter(temp_path, engine="openpyxl") as writer:
            pd.DataFrame(summary_rows, columns=["Metric", "Value"]).to_excel(
                writer, sheet_name="Summary", index=False
            )
            issues_df.to_excel(writer, sheet_name="Validation Issues", index=False)
            changes_df.to_excel(writer, sheet_name="Transform Changes", index=False)
            audit_df.to_excel(writer, sheet_name="Audit Log", index=False)

        _autosize_excel_columns(temp_path)
        os.replace(temp_path, output_path)
    except PermissionError as exc:
        _raise_excel_permission_error(output_path, exc)
    finally:
        temp_path.unlink(missing_ok=True)

    return output_path


def _add_key_value_table(document: Document, heading: str, rows: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Value"

    for metric, value in rows:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = value


def _add_issues_table(document: Document, issues: list[dict]) -> None:
    document.add_heading("Validation Issue Details", level=1)

    if not issues:
        document.add_paragraph("No validation issues were found.")
        return

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Severity"
    header_cells[1].text = "Column"
    header_cells[2].text = "Message"

    for issue in issues:
        cells = table.add_row().cells
        cells[0].text = str(issue["severity"])
        cells[1].text = str(issue["column"])
        cells[2].text = str(issue["message"])


def _add_list(document: Document, heading: str, entries: list[str]) -> None:
    document.add_heading(heading, level=1)

    if not entries:
        document.add_paragraph("No entries.")
        return

    for entry in entries:
        document.add_paragraph(str(entry), style="List Bullet")


def _autosize_excel_columns(output_path: Path) -> None:
    from openpyxl import load_workbook

    workbook = load_workbook(output_path)

    for worksheet in workbook.worksheets:
        for column_cells in worksheet.columns:
            column_letter = column_cells[0].column_letter
            max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
            worksheet.column_dimensions[column_letter].width = min(max(max_length + 2, 12), 80)

    workbook.save(output_path)


def _raise_excel_permission_error(output_path: Path, exc: PermissionError) -> None:
    lock_path = output_path.with_name(f"~${output_path.name}")
    lock_hint = (
        f" Excel lock file detected at {lock_path}."
        if lock_path.exists()
        else ""
    )
    raise RuntimeError(
        f"Could not write Excel report to {output_path}. "
        "Close the workbook if it is open in Excel, then run the pipeline again."
        f"{lock_hint}"
    ) from exc
